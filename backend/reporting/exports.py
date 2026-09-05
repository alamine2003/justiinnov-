"""Génération des exports Excel, CSV, Word et PDF (§5.7).

L'export des dépenses reprend délibérément les colonnes du fichier Excel
d'origine — N°ORDRE, DATE, TEAM, OWNER, LIBELLE DES TRANSACTIONS, DEPENSES,
MONTANT JUSTIFIER, ECART, PIECES JUSTIFICATIVES — pour que les rapprochements
avec l'historique restent possibles.

Un même jeu de lignes (:class:`Tableau`) alimente les trois formats
tabulaires : le classeur Excel, le CSV et le document Word. Les en-têtes de
colonnes sont le **contrat de fichier** et restent en français quel que soit
le format ; seuls les intitulés de document (titres, en-tête de période) et
le PDF suivent la langue de l'utilisateur.
"""

import csv
from dataclasses import dataclass, field
from decimal import Decimal
from io import BytesIO, StringIO

from django.utils import timezone
from django.utils.translation import gettext as _
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from budget.aggregates import budget_figures
from expenses.models import Proof

from .scope import fuseau_de

ZERO = Decimal("0.00")

#: Dossiers listés au plus dans le rapport PDF, qui reste une synthèse.
MAX_DOSSIERS_PDF = 200

HEADER_FILL = PatternFill("solid", fgColor="1F3864")
HEADER_FONT = Font(color="FFFFFF", bold=True)

#: Types MIME des formats produits.
XLSX = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
CSV = "text/csv; charset=utf-8"
PDF = "application/pdf"

EXPENSE_COLUMNS = [
    ("N°ORDRE", 16),
    ("DATE", 18),
    ("PAYS", 16),
    ("TEAM", 18),
    ("OWNER", 20),
    ("LIBELLE DES TRANSACTIONS", 38),
    ("DEPENSES", 16),
    # Ce que porte la pièce, quand le décaissement a eu lieu dans une autre
    # devise : sans ces colonnes, le rapprochement avec le justificatif est
    # impossible, aucun des chiffres du classeur n'y figurant.
    ("DEVISE D'ORIGINE", 16),
    ("MONTANT D'ORIGINE", 18),
    ("MONTANT JUSTIFIER", 18),
    ("ECART", 14),
    ("STATUT", 14),
    ("PIECES JUSTIFICATIVES", 30),
]

RECONCILIATION_COLUMNS = [
    ("PAYS", 20),
    ("ENVELOPPE", 30),
    ("DEVISE", 10),
    ("BUDGET", 16),
    ("ENGAGE", 16),
    ("CONSOMME", 16),
    ("JUSTIFIE", 16),
    ("ECART", 16),
    ("DISPONIBLE", 16),
    ("TAUX JUSTIFICATION", 20),
]

DOSSIER_COLUMNS = [
    ("N°ORDRE", 16), ("LIBELLE", 34), ("PAYS", 18), ("DATE", 14),
    ("STATUT", 14), ("DEPENSES", 16), ("JUSTIFIE", 16), ("ECART", 16),
    ("PIECES", 12),
]


@dataclass
class Tableau:
    """Lignes d'un export, indépendantes du format qui les écrira.

    ``titre`` nomme la feuille Excel, la section du CSV et le sous-titre
    Word. ``total`` est une ligne de même largeur que les autres, ``None``
    dans les colonnes sans total — ou ``None`` tout court quand un total
    n'aurait pas de sens.
    """

    titre: str
    colonnes: list
    lignes: list = field(default_factory=list)
    total: list | None = None


#: Premiers caractères qu'un tableur interprète comme une formule.
FORMULE = ("=", "+", "-", "@", "\t", "\r")


def cellule_sure(value):
    """Texte inoffensif pour un tableur.

    Un libellé saisi par un utilisateur peut commencer par ``=`` : Excel y
    verrait une formule, ``=HYPERLINK(...)`` ou un appel DDE, exécutée à
    l'ouverture du classeur sur le poste de qui l'ouvre. La valeur est
    conservée telle quelle — le rapprochement doit retrouver le libellé — mais
    marquée comme texte à l'écriture (voir :func:`ecrire`).
    """
    return isinstance(value, str) and value.startswith(FORMULE)


def ecrire(sheet, row, column, value):
    """Écrit une cellule en forçant le texte quand il ressemble à une formule.

    Les montants sont écrits en ``Decimal``, sans passer par ``float`` : un
    arrondi binaire ferait apparaître 0,1 + 0,2 ≠ 0,3 dans un rapport dont
    la raison d'être est l'exactitude des écarts.
    """
    cell = sheet.cell(row=row, column=column, value=value)
    if cellule_sure(value):
        cell.data_type = "s"
    return cell


def _style_header(sheet, columns):
    for index, (title, width) in enumerate(columns, start=1):
        cell = sheet.cell(row=1, column=index, value=title)
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        sheet.column_dimensions[get_column_letter(index)].width = width
    sheet.freeze_panes = "A2"


def _proof_summary(dossier):
    """Reprend la nuance « Reçu (justif incomplet) » du fichier d'origine."""
    labels = []
    for proof in dossier.proofs.all():
        if proof.status == Proof.ProofStatus.ARCHIVED:
            continue
        label = proof.get_kind_display()
        if not proof.is_complete:
            label += _(" (justif incomplet)")
        labels.append(label)
    return " ; ".join(labels)


def _total_si_devise_unique(devises, ligne):
    """La ligne de total, ou ``None`` quand les devises se mélangent.

    Additionner des francs CFA et des francs guinéens donne un chiffre sans
    unité, présenté comme un total : le tableau de bord s'y refuse, l'export
    aussi. Un export vide n'a pas de total non plus.
    """
    return ligne if len(devises) == 1 else None


# --- Lignes ------------------------------------------------------------------


def lignes_depenses(expenses):
    """Lignes de dépenses, classées par leur propre date, au format historique.

    Une ligne se classe par sa date, pas par celle du dossier qui la porte :
    un dossier ouvert en décembre peut recevoir une ligne de janvier, et
    cette ligne relève de l'exercice suivant — comme pour son imputation.
    L'heure affichée est celle du pays de la ligne : c'est celle qu'on lira
    sur la pièce.
    """
    tableau = Tableau("BASE DE DONNEES ACTIONS", EXPENSE_COLUMNS)
    source = (
        expenses.select_related("dossier__country", "country", "team", "owner")
        .prefetch_related("dossier__proofs")
        .order_by("date", "pk")
    )
    totals = {"amount": ZERO, "justified": ZERO}
    devises = set()
    # Le résumé des pièces est calculé une fois par dossier, pas par ligne.
    pieces = {}

    for expense in source:
        dossier = expense.dossier
        if dossier.pk not in pieces:
            pieces[dossier.pk] = _proof_summary(dossier)
        totals["amount"] += expense.amount
        totals["justified"] += expense.justified_amount
        devises.add(expense.country.currency)
        tableau.lignes.append([
            dossier.number,
            timezone.localtime(expense.date, fuseau_de(expense.country)).strftime(
                "%d/%m/%Y %H:%M"
            ),
            expense.country.name,
            expense.team.name if expense.team else "",
            expense.owner.name if expense.owner else "",
            expense.title,
            expense.amount,
            expense.original_currency or "",
            expense.original_amount
            if expense.original_amount is not None
            else "",
            expense.justified_amount,
            expense.amount - expense.justified_amount,
            expense.get_status_display(),
            pieces[dossier.pk],
        ])

    tableau.total = _total_si_devise_unique(devises, [
        None, None, None, None, None, "TOTAL",
        totals["amount"], None, None, totals["justified"],
        totals["amount"] - totals["justified"], None, None,
    ])
    return tableau


def tableaux_rapprochement(budgets, dossiers):
    """Rapprochement enveloppe par enveloppe, puis dossier par dossier."""
    enveloppes = Tableau("Rapprochement budgets", RECONCILIATION_COLUMNS)
    totaux = {k: ZERO for k in ("amount", "engaged", "consumed", "justified", "gap", "remaining")}
    devises = set()
    for budget in budgets:
        figures = budget_figures(budget)
        rate = figures["justification_rate"]
        devises.add(budget.country.currency)
        totaux["amount"] += budget.amount
        for cle in ("engaged", "consumed", "justified", "gap", "remaining"):
            totaux[cle] += figures[cle]
        enveloppes.lignes.append([
            budget.country.name,
            # ``scope_label`` distingue projet, équipe et manager : une
            # sous-enveloppe d'équipe se lisait « Enveloppe du pays ».
            budget.scope_label or _("Enveloppe du pays"),
            budget.country.currency,
            budget.amount,
            figures["engaged"],
            figures["consumed"],
            figures["justified"],
            figures["gap"],
            figures["remaining"],
            rate if rate is not None else "",
        ])
    enveloppes.total = _total_si_devise_unique(devises, [
        "TOTAL", None, next(iter(devises), None),
        totaux["amount"], totaux["engaged"], totaux["consumed"],
        totaux["justified"], totaux["gap"], totaux["remaining"], None,
    ])

    detail = Tableau("Rapprochement dossiers", DOSSIER_COLUMNS)
    totaux = {"amount": ZERO, "justified": ZERO, "gap": ZERO}
    devises = set()
    for dossier in dossiers.select_related("country"):
        totals = dossier.totals()
        devises.add(dossier.country.currency)
        for cle in totaux:
            totaux[cle] += totals[cle]
        detail.lignes.append([
            dossier.number,
            dossier.label,
            dossier.country.name,
            dossier.date.strftime("%d/%m/%Y"),
            dossier.get_status_display(),
            totals["amount"],
            totals["justified"],
            totals["gap"],
            dossier.counts()["proofs"],
        ])
    # Le total se lit dans LIBELLE, N°ORDRE restant vide : comme dans le
    # classeur des dépenses, la colonne des numéros ne porte que des numéros.
    detail.total = _total_si_devise_unique(devises, [
        None, "TOTAL", None, None, None,
        totaux["amount"], totaux["justified"], totaux["gap"], None,
    ])
    return [enveloppes, detail]


# --- Écritures ---------------------------------------------------------------


def classeur_xlsx(tableaux, contexte=None):
    """Un classeur, une feuille par tableau, l'en-tête figé et stylé."""
    workbook = Workbook()
    for index, tableau in enumerate(tableaux):
        sheet = workbook.active if index == 0 else workbook.create_sheet()
        sheet.title = tableau.titre
        _style_header(sheet, tableau.colonnes)
        row_index = 2
        for ligne in tableau.lignes:
            for column, value in enumerate(ligne, start=1):
                ecrire(sheet, row_index, column, value)
            row_index += 1
        if tableau.total:
            for column, value in enumerate(tableau.total, start=1):
                # Une cellule sans total reste vide, pas « None » ni « » :
                # l'import relit ce classeur et reconnaît la ligne TOTAL à
                # son N°ORDRE vide.
                if value is not None:
                    ecrire(sheet, row_index, column, value).font = Font(bold=True)
    return _to_bytes(workbook)


def _cellule_csv(value):
    """Valeur telle qu'Excel francophone la lira.

    Les montants gardent la virgule décimale ; un texte en forme de formule
    reçoit une apostrophe en tête — un CSV n'a pas de type de cellule, c'est
    la seule façon de dire à un tableur « ceci est du texte ».
    """
    if value is None:
        return ""
    if isinstance(value, Decimal):
        return str(value).replace(".", ",")
    if cellule_sure(value):
        return "'" + value
    return str(value)


def fichier_csv(tableaux, contexte=None):
    """CSV pour Excel francophone : UTF-8 avec BOM, séparateur « ; ».

    Sans la marque d'ordre, Excel lit le fichier en Latin-1 et les accents
    se brisent ; sans le point-virgule, il ne sépare pas les colonnes. Un
    CSV n'a pas de feuilles : les tableaux se suivent, chacun sous son titre.
    """
    buffer = StringIO()
    writer = csv.writer(buffer, delimiter=";", lineterminator="\r\n")
    for index, tableau in enumerate(tableaux):
        if index:
            writer.writerow([])
            writer.writerow([tableau.titre])
        writer.writerow([title for title, _largeur in tableau.colonnes])
        for ligne in tableau.lignes:
            writer.writerow([_cellule_csv(value) for value in ligne])
        if tableau.total:
            writer.writerow([_cellule_csv(value) for value in tableau.total])
    return ("\ufeff" + buffer.getvalue()).encode("utf-8")


def _cellule_docx(value):
    if value is None:
        return ""
    if isinstance(value, Decimal):
        return _fmt(value)
    return str(value)


def document_docx(tableaux, contexte):
    """Document Word : titre, en-tête de période, un tableau par section.

    ``contexte`` porte ``titre``, ``pays`` et ``periode``, déjà dans la
    langue de l'utilisateur.
    """
    document = Document()
    section = document.sections[0]
    # Paysage : treize colonnes ne tiennent pas dans la largeur d'un portrait.
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    document.styles["Normal"].font.size = Pt(9)

    document.add_heading(contexte["titre"], level=1)
    document.add_paragraph(
        _("Pays : %(pays)s · Exercice : %(exercice)s · Période : %(periode)s")
        % contexte
    )
    document.add_paragraph(_("Montants exprimés dans la devise de chaque pays."))

    for tableau in tableaux:
        document.add_heading(tableau.titre, level=2)
        if not tableau.lignes:
            document.add_paragraph(_("Aucune donnée sur la période."))
            continue
        table = document.add_table(rows=1, cols=len(tableau.colonnes))
        table.style = "Table Grid"
        for cell, (title, _largeur) in zip(table.rows[0].cells, tableau.colonnes):
            cell.text = ""
            cell.paragraphs[0].add_run(title).bold = True
        for ligne in tableau.lignes:
            for cell, value in zip(table.add_row().cells, ligne):
                cell.text = _cellule_docx(value)
        if tableau.total:
            for cell, value in zip(table.add_row().cells, tableau.total):
                cell.text = ""
                cell.paragraphs[0].add_run(_cellule_docx(value)).bold = True

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


#: Écriture, type MIME et extension de chaque format tabulaire.
FORMATS = {
    "xlsx": (classeur_xlsx, XLSX),
    "csv": (fichier_csv, CSV),
    "docx": (document_docx, DOCX),
}


def build_expenses_workbook(expenses):
    """Classeur des lignes de dépenses, classées par date."""
    return classeur_xlsx([lignes_depenses(expenses)])


def build_reconciliation_workbook(budgets, dossiers):
    """Classeur de rapprochement, enveloppes puis dossiers."""
    return classeur_xlsx(tableaux_rapprochement(budgets, dossiers))


def _to_bytes(workbook):
    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _fmt(value):
    return f"{Decimal(value):,.2f}".replace(",", " ")


# --- PDF ---------------------------------------------------------------------


def build_country_report_pdf(budgets, dossiers, expenses, periode):
    """Rapport de synthèse par pays, avec ses enveloppes et ses dossiers."""
    libelle = periode.libelle
    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        leftMargin=15 * mm, rightMargin=15 * mm,
        topMargin=15 * mm, bottomMargin=15 * mm,
        title=_("Rapport de contrôle budgétaire — %(periode)s") % {"periode": libelle},
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitreRapport", parent=styles["Title"], fontSize=18, spaceAfter=4
    )
    story = [
        Paragraph(_("Contrôle budgétaire — %(periode)s") % {"periode": libelle}, title_style),
        Paragraph(_("Montants exprimés dans la devise de chaque pays."), styles["Normal"]),
        Spacer(1, 8 * mm),
    ]

    budget_rows = [[
        _("Pays"), _("Enveloppe"), _("Budget"), _("Engagé"), _("Consommé"),
        _("Justifié"), _("Disponible"),
    ]]
    for budget in budgets:
        figures = budget_figures(budget)
        budget_rows.append([
            budget.country.name,
            budget.scope_label or _("Enveloppe du pays"),
            _fmt(budget.amount),
            _fmt(figures["engaged"]),
            _fmt(figures["consumed"]),
            _fmt(figures["justified"]),
            _fmt(figures["remaining"]),
        ])

    story.append(Paragraph(_("Enveloppes"), styles["Heading2"]))
    story.append(_table(budget_rows) if len(budget_rows) > 1 else _empty(styles))
    story.append(Spacer(1, 8 * mm))

    dossier_rows = [[
        "N°ORDRE", _("Libellé"), _("Pays"), _("Date"), _("Statut"), _("Dépenses"),
        _("Justifié"), _("Écart"),
    ]]
    # Le PDF est une synthèse : au-delà de MAX_DOSSIERS_PDF dossiers, le
    # lecteur est renvoyé au classeur Excel, et le document le dit — un
    # rapport tronqué en silence ferait croire à un exercice plus court.
    nombre_dossiers = dossiers.count()
    for dossier in dossiers.select_related("country")[:MAX_DOSSIERS_PDF]:
        totals = dossier.totals()
        dossier_rows.append([
            dossier.number,
            dossier.label[:45],
            dossier.country.name,
            dossier.date.strftime("%d/%m/%Y"),
            dossier.get_status_display(),
            _fmt(totals["amount"]),
            _fmt(totals["justified"]),
            _fmt(totals["gap"]),
        ])

    story.append(Paragraph(_("Dossiers de justification"), styles["Heading2"]))
    troncature = avertissement_troncature(nombre_dossiers)
    if troncature:
        story.append(Paragraph(troncature, styles["Italic"]))
    story.append(_table(dossier_rows) if len(dossier_rows) > 1 else _empty(styles))

    document.build(story)
    return buffer.getvalue()


def avertissement_troncature(nombre_dossiers):
    """Phrase qui signale les dossiers absents du PDF, ou ``None``."""
    if nombre_dossiers <= MAX_DOSSIERS_PDF:
        return None
    return _(
        "Seuls les %(max)s dossiers les plus récents sur %(total)s figurent "
        "ici ; la liste complète est dans l'export Excel des dépenses."
    ) % {"max": MAX_DOSSIERS_PDF, "total": nombre_dossiers}


def _empty(styles):
    return Paragraph(_("Aucune donnée sur la période."), styles["Italic"])


def _table(rows):
    table = Table(rows, repeatRows=1, hAlign="LEFT")
    table.setStyle(
        TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F3864")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("ALIGN", (2, 1), (-1, -1), "RIGHT"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B0B7C3")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F4F8")]),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
        ])
    )
    return table
