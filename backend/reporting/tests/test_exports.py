"""Exports : cellules sûres, montants exacts, périmètre et trace."""

from datetime import date, datetime
from decimal import Decimal
from io import BytesIO
from unittest import mock

from django.utils import timezone, translation
from docx import Document
from openpyxl import Workbook, load_workbook
from rest_framework import status

from accounts.models import Role
from accounts.tests.test_scoping import make_user
from budget.models import Budget
from expenses.models import AuditLog, Dossier, Expense
from expenses.tests.base import in_memory_storage
from expenses.workflow import Status
from reporting import exports
from reporting.scope import Periode
from reporting.tests.test_dashboard import DashboardTestCase

#: Toutes les routes d'export, avec leur type MIME.
ROUTES = {
    "expenses.xlsx": exports.XLSX,
    "expenses.csv": exports.CSV,
    "expenses.docx": exports.DOCX,
    "reconciliation.xlsx": exports.XLSX,
    "reconciliation.csv": exports.CSV,
    "reconciliation.docx": exports.DOCX,
    "report.pdf": exports.PDF,
}


def _colonne(sheet, nom):
    entetes = [cell.value for cell in sheet[1]]
    index = entetes.index(nom) + 1
    return [sheet.cell(row=r, column=index) for r in range(2, sheet.max_row + 1)]


class CelluleSureTests(DashboardTestCase):
    def test_un_libelle_en_forme_de_formule_est_ecrit_en_texte(self):
        """Un libellé « =HYPERLINK(...) » s'exécuterait à l'ouverture du
        classeur sur le poste du contrôleur."""
        piege = '=HYPERLINK("http://exemple.test";"cliquez")'
        self.make_expense(title=piege, status=Status.JUSTIFIED)
        self.login(self.doo)

        response = self.client.get("/api/exports/expenses.xlsx", {"year": self.year})

        sheet = load_workbook(BytesIO(response.content)).active
        cellule = next(c for c in _colonne(sheet, "LIBELLE DES TRANSACTIONS") if c.value == piege)
        self.assertEqual(cellule.data_type, "s")

    def test_chaque_prefixe_dangereux_est_neutralise(self):
        sheet = Workbook().active
        for prefixe in ("=", "+", "-", "@", "\t", "\r"):
            with self.subTest(prefixe=repr(prefixe)):
                cellule = exports.ecrire(sheet, 1, 1, prefixe + "SUM(A1)")
                self.assertEqual(cellule.data_type, "s")
        self.assertEqual(exports.ecrire(sheet, 1, 2, "Carburant").data_type, "s")
        self.assertFalse(exports.cellule_sure(Decimal("1.00")))

    def test_le_libelle_du_dossier_est_aussi_protege(self):
        self.dossier.label = "@SUM(1+1)"
        self.dossier.save()
        self.login(self.doo)

        response = self.client.get("/api/exports/reconciliation.xlsx", {"year": self.year})

        detail = load_workbook(BytesIO(response.content))["Rapprochement dossiers"]
        cellule = next(c for c in _colonne(detail, "LIBELLE") if c.value == "@SUM(1+1)")
        self.assertEqual(cellule.data_type, "s")


class MontantsExactsTests(DashboardTestCase):
    def test_les_montants_sont_ecrits_en_decimal(self):
        """Passer par ``float`` introduit des arrondis binaires dans un
        rapport dont la raison d'être est l'exactitude des écarts."""
        sheet = Workbook().active

        cellule = exports.ecrire(sheet, 1, 1, Decimal("12345678901234.56"))

        self.assertIsInstance(cellule.value, Decimal)
        self.assertEqual(cellule.data_type, "n")

    def test_le_classeur_relu_porte_le_montant_exact(self):
        self.login(self.doo)

        response = self.client.get("/api/exports/expenses.xlsx", {"year": self.year})

        sheet = load_workbook(BytesIO(response.content)).active
        montants = {c.value for c in _colonne(sheet, "DEPENSES")}
        self.assertIn(300000, montants)


class SousEnveloppesTests(DashboardTestCase):
    def test_une_sous_enveloppe_d_equipe_est_libellee(self):
        Budget.objects.create(
            country=self.togo, year=self.year, team=self.team,
            amount=Decimal("100000.00"),
        )
        self.login(self.doo)

        response = self.client.get("/api/exports/reconciliation.xlsx", {"year": self.year})

        sheet = load_workbook(BytesIO(response.content))["Rapprochement budgets"]
        libelles = {c.value for c in _colonne(sheet, "ENVELOPPE")}
        self.assertIn("Équipe Équipe Lomé", libelles)
        self.assertIn("Enveloppe du pays", libelles)


class PerimetreEtTraceTests(DashboardTestCase):
    def test_un_pays_inconnu_repond_404_sans_trace(self):
        """Régression : l'audit de l'export échouait sur une clé étrangère
        inexistante, en erreur 500."""
        self.login(self.doo)

        response = self.client.get(
            "/api/exports/expenses.xlsx", {"year": self.year, "country": 999999}
        )

        self.assertEqual(response.status_code, status.HTTP_404_NOT_FOUND)
        self.assertFalse(AuditLog.objects.filter(object_type="Export").exists())

    def test_un_pays_hors_perimetre_repond_404(self):
        """Les administrateurs voient tout : le cas ne se présente plus par
        l'API, mais le cloisonnement du queryset reste vérifié, comme pour
        le tableau de bord."""
        from rest_framework.exceptions import NotFound

        from accounts.permissions import get_access
        from reporting.scope import querysets_pour

        with self.assertRaises(NotFound):
            querysets_pour(get_access(self.rep_ivoire), self.year, self.togo.pk)

    def test_une_annee_invalide_repond_400(self):
        self.login(self.doo)

        response = self.client.get("/api/exports/expenses.xlsx", {"year": "2024;rm"})

        self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)

    def test_le_nom_du_fichier_reprend_l_annee_validee(self):
        self.login(self.doo)

        response = self.client.get("/api/exports/expenses.xlsx", {"year": "02024"})

        self.assertEqual(
            response["Content-Disposition"], 'attachment; filename="depenses-2024.xlsx"'
        )

    def test_la_trace_porte_l_adresse_du_client(self):
        self.login(self.doo)

        self.client.get(
            "/api/exports/report.pdf", {"year": self.year},
            REMOTE_ADDR="10.20.30.40", HTTP_X_FORWARDED_FOR="1.2.3.4",
        )

        entry = AuditLog.objects.get(object_type="Export")
        self.assertEqual(entry.ip_address, "10.20.30.40")
        self.assertEqual(entry.detail["year"], self.year)
        self.assertEqual(entry.detail["format"], "pdf")
        self.assertIsNone(entry.detail["month"])


@in_memory_storage
class RapportPdfTests(DashboardTestCase):
    def test_la_troncature_est_annoncee(self):
        self.assertIsNone(exports.avertissement_troncature(exports.MAX_DOSSIERS_PDF))
        with mock.patch.object(exports, "MAX_DOSSIERS_PDF", 2):
            phrase = exports.avertissement_troncature(5)
        self.assertIn("2 dossiers", phrase)
        self.assertIn("sur 5", phrase)

    def test_le_rapport_se_construit_avec_l_avertissement(self):
        for index in range(3):
            Dossier.objects.create(
                number=f"P-{index}", label=f"Dossier {index}", country=self.togo,
                date=date(self.year, 2, 1), status=Status.SUBMITTED,
            )
        self.login(self.doo)

        with mock.patch.object(exports, "MAX_DOSSIERS_PDF", 2), mock.patch.object(
            exports, "avertissement_troncature", wraps=exports.avertissement_troncature
        ) as avertissement:
            response = self.client.get("/api/exports/report.pdf", {"year": self.year})

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertTrue(response.content.startswith(b"%PDF"))
        avertissement.assert_called_once_with(4)


class AccesTests(DashboardTestCase):
    """Exports réservés aux administrateurs : lecture comprise."""

    def test_les_autres_roles_recoivent_403_sans_trace(self):
        for user in (self.controller, self.rep_ivoire, self.owner):
            self.login(user)
            for route in ROUTES:
                with self.subTest(role=user.profile.role, route=route):
                    response = self.client.get(f"/api/exports/{route}", {"year": self.year})
                    self.assertEqual(response.status_code, status.HTTP_403_FORBIDDEN)
        self.assertFalse(AuditLog.objects.filter(object_type="Export").exists())

    def test_les_administrateurs_exportent_dans_tous_les_formats(self):
        rh = make_user("rh.admin", Role.ADMIN)
        for user in (self.doo, rh):
            self.login(user)
            for route, content_type in ROUTES.items():
                with self.subTest(role=user.profile.role, route=route):
                    response = self.client.get(f"/api/exports/{route}", {"year": self.year})
                    self.assertEqual(response.status_code, status.HTTP_200_OK)
                    self.assertEqual(response["Content-Type"], content_type)
        self.assertEqual(
            AuditLog.objects.filter(object_type="Export").count(), 2 * len(ROUTES)
        )


class FormatsTests(DashboardTestCase):
    def _export(self, route, **params):
        self.login(self.doo)
        return self.client.get(f"/api/exports/{route}", {"year": self.year, **params})

    def test_le_csv_est_lisible_par_excel_francophone(self):
        """BOM, point-virgule, virgule décimale : sans cela Excel brise les
        accents et ne sépare pas les colonnes."""
        piege = '=HYPERLINK("http://exemple.test";"cliquez")'
        self.make_expense(title=piege, status=Status.JUSTIFIED)

        response = self._export("expenses.csv")

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(response["Content-Type"], exports.CSV)
        self.assertEqual(
            response["Content-Disposition"],
            f'attachment; filename="depenses-{self.year}.csv"',
        )
        self.assertTrue(response.content.startswith(b"\xef\xbb\xbf"))
        texte = response.content.decode("utf-8-sig")
        lignes = texte.splitlines()
        self.assertEqual(lignes[0].split(";"), [t for t, _ in exports.EXPENSE_COLUMNS])
        self.assertIn(";300000,00;", texte)
        # La formule est neutralisée par une apostrophe : un CSV n'a pas de
        # type de cellule.
        self.assertIn("'=HYPERLINK", texte)
        self.assertNotIn(";=HYPERLINK", texte)
        self.assertIn(";TOTAL;", lignes[-1])

    def test_le_word_reprend_les_colonnes_de_l_excel_avec_les_totaux(self):
        response = self._export("expenses.docx", country=self.togo.pk)

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertEqual(response["Content-Type"], exports.DOCX)
        self.assertEqual(
            response["Content-Disposition"],
            f'attachment; filename="depenses-{self.year}.docx"',
        )
        document = Document(BytesIO(response.content))
        table = document.tables[0]
        self.assertEqual(
            [c.text for c in table.rows[0].cells], [t for t, _ in exports.EXPENSE_COLUMNS]
        )
        total = [c.text for c in table.rows[-1].cells]
        self.assertEqual(total[5], "TOTAL")
        self.assertEqual(total[6], "500 000.00")
        self.assertEqual(total[9], "250 000.00")
        entete = "\n".join(p.text for p in document.paragraphs)
        self.assertIn("Pays : Togo", entete)
        self.assertIn(f"Exercice : {self.year}", entete)
        self.assertIn(f"Période : exercice {self.year}", entete)

    def test_le_rapprochement_existe_en_csv_et_en_word(self):
        csv_ = self._export("reconciliation.csv")
        docx = self._export("reconciliation.docx")

        texte = csv_.content.decode("utf-8-sig")
        self.assertIn(";".join(t for t, _ in exports.RECONCILIATION_COLUMNS), texte)
        self.assertIn("Rapprochement dossiers", texte)
        self.assertIn(";".join(t for t, _ in exports.DOSSIER_COLUMNS), texte)
        document = Document(BytesIO(docx.content))
        self.assertEqual(len(document.tables), 2)
        self.assertEqual(
            [c.text for c in document.tables[0].rows[0].cells],
            [t for t, _ in exports.RECONCILIATION_COLUMNS],
        )
        self.assertIn("Tous les pays", "\n".join(p.text for p in document.paragraphs))

    def test_le_classeur_porte_les_totaux(self):
        response = self._export("reconciliation.xlsx")

        workbook = load_workbook(BytesIO(response.content))
        enveloppes = list(workbook["Rapprochement budgets"].iter_rows(values_only=True))
        self.assertEqual(enveloppes[-1][0], "TOTAL")
        self.assertEqual(enveloppes[-1][3], 1500000)
        dossiers = list(workbook["Rapprochement dossiers"].iter_rows(values_only=True))
        self.assertIsNone(dossiers[-1][0])
        self.assertEqual(dossiers[-1][1], "TOTAL")
        self.assertEqual(dossiers[-1][5], 500000)

    def test_les_totaux_ne_melangent_pas_les_devises(self):
        """Additionner des francs CFA et des francs guinéens donnerait un
        chiffre sans unité : le total est tu, comme au tableau de bord."""
        self.ivoire.currency = "GNF"
        self.ivoire.save()
        abidjan = Dossier.objects.create(
            number="CI-0001", label="Salon", country=self.ivoire,
            date=date(self.year, 4, 2), status=Status.SUBMITTED,
        )
        self.make_expense(
            dossier=abidjan, country=self.ivoire, team=None, owner=None,
            status=Status.SUBMITTED, budget=self.budget_ivoire,
        )

        self.assertIsNone(exports.lignes_depenses(Expense.objects.all()).total)
        self.assertIsNotNone(
            exports.lignes_depenses(Expense.objects.filter(country=self.togo)).total
        )
        enveloppes, dossiers = exports.tableaux_rapprochement(
            Budget.objects.with_consumption(), Dossier.objects.with_totals()
        )
        self.assertIsNone(enveloppes.total)
        self.assertIsNone(dossiers.total)

    def test_un_export_vide_n_a_pas_de_total(self):
        self.assertIsNone(exports.lignes_depenses(Expense.objects.none()).total)


class PeriodeTests(DashboardTestCase):
    """Classement par mois : la date du dossier pour le rapprochement, celle
    de chaque ligne pour l'export des dépenses."""

    @classmethod
    def setUpTestData(cls):
        super().setUpTestData()
        # Les lignes du socle sont datées d'aujourd'hui : elles rejoignent
        # le mois de leur dossier, mars, pour que l'export par mois les voie.
        Expense.objects.filter(dossier=cls.dossier).update(
            date=timezone.make_aware(datetime(cls.year, 3, 15, 10))
        )
        cls.juillet = Dossier.objects.create(
            number="N-JUIL", label="Tournée", country=cls.togo,
            date=date(cls.year, 7, 10), status=Status.SUBMITTED,
        )
        cls.make_expense(
            cls, dossier=cls.juillet, status=Status.SUBMITTED, budget=cls.budget,
            date=timezone.make_aware(datetime(cls.year, 7, 10, 9)),
        )

    def _export(self, route, **params):
        self.login(self.doo)
        return self.client.get(f"/api/exports/{route}", {"year": self.year, **params})

    def _numeros_csv(self, response):
        lignes = response.content.decode("utf-8-sig").splitlines()[1:]
        return {ligne.split(";")[0] for ligne in lignes} - {""}

    def test_sans_mois_l_exercice_entier(self):
        self.assertEqual(self._numeros_csv(self._export("expenses.csv")), {"N-0001", "N-JUIL"})

    def test_le_mois_borne_les_dossiers_et_nomme_le_fichier(self):
        mars = self._export("expenses.csv", month=3)
        juillet = self._export("expenses.csv", month="7")

        self.assertEqual(self._numeros_csv(mars), {"N-0001"})
        self.assertEqual(self._numeros_csv(juillet), {"N-JUIL"})
        self.assertEqual(
            mars["Content-Disposition"],
            f'attachment; filename="depenses-{self.year}-03.csv"',
        )

    def test_le_rapprochement_garde_l_enveloppe_annuelle(self):
        """Une enveloppe est annuelle : le mois ne filtre que les dossiers."""
        response = self._export("reconciliation.xlsx", month=7)

        workbook = load_workbook(BytesIO(response.content))
        enveloppes = [r[1] for r in workbook["Rapprochement budgets"].iter_rows(min_row=2, values_only=True)]
        dossiers = {r[0] for r in workbook["Rapprochement dossiers"].iter_rows(min_row=2, values_only=True)}
        self.assertIn("Enveloppe du pays", enveloppes)
        self.assertEqual(dossiers - {None}, {"N-JUIL"})

    def test_le_pdf_accepte_le_mois(self):
        response = self._export("report.pdf", month=7)

        self.assertEqual(response.status_code, status.HTTP_200_OK)
        self.assertTrue(response.content.startswith(b"%PDF"))
        self.assertEqual(
            response["Content-Disposition"],
            f'attachment; filename="rapport-{self.year}-07.pdf"',
        )

    def test_un_mois_invalide_repond_400(self):
        for mois in ("0", "13", "mars"):
            with self.subTest(mois=mois):
                response = self._export("expenses.xlsx", month=mois)
                self.assertEqual(response.status_code, status.HTTP_400_BAD_REQUEST)
                self.assertIn("month", response.data)
        self.assertFalse(AuditLog.objects.filter(object_type="Export").exists())

    def test_la_trace_porte_le_format_et_le_mois(self):
        self._export("expenses.docx", month=3, country=self.togo.pk)

        entry = AuditLog.objects.get(object_type="Export")
        self.assertEqual(entry.detail["format"], "docx")
        self.assertEqual(entry.detail["month"], 3)
        self.assertEqual(entry.detail["country"], self.togo.pk)
        self.assertIn(f"depenses-{self.year}-03.docx", entry.label)

    def test_le_libelle_de_periode_suit_la_langue(self):
        self.assertEqual(Periode(self.year, 3).libelle, f"mars {self.year}")
        self.assertEqual(Periode(self.year).libelle, f"exercice {self.year}")
        with translation.override("en"):
            self.assertEqual(Periode(self.year, 3).libelle, f"March {self.year}")
            self.assertEqual(Periode(self.year).libelle, f"fiscal year {self.year}")
        self.assertEqual(Periode(self.year, 3).suffixe, f"{self.year}-03")

    def test_l_en_tete_word_nomme_le_mois(self):
        response = self._export("expenses.docx", month=3)

        document = Document(BytesIO(response.content))
        self.assertIn(
            f"Période : mars {self.year}", "\n".join(p.text for p in document.paragraphs)
        )


class LangueDesTitresTests(DashboardTestCase):
    def test_le_titre_du_document_word_suit_la_langue(self):
        self.login(self.doo)

        with translation.override("en"):
            response = self.client.get(
                "/api/exports/expenses.docx", {"year": self.year},
                HTTP_ACCEPT_LANGUAGE="en",
            )

        document = Document(BytesIO(response.content))
        titre = document.paragraphs[0].text
        self.assertIn("Expense export", titre)
        self.assertIn(f"fiscal year {self.year}", titre)
        # Le journal, lui, reste en français : il se relit tel qu'il a été écrit.
        self.assertTrue(
            AuditLog.objects.filter(label__startswith="Export des dépenses").exists()
        )

    def test_la_mention_de_piece_incomplete_suit_la_langue(self):
        from expenses.models import Proof

        Proof.objects.create(
            dossier=self.dossier, file="justificatifs/test.pdf",
            original_name="facture.pdf", kind=Proof.Kind.INVOICE,
            is_complete=False, sha256="a" * 64,
        )

        with translation.override("en"):
            tableau = exports.lignes_depenses(Expense.objects.filter(country=self.togo))

        self.assertIn("(incomplete proof)", tableau.lignes[0][-1])
